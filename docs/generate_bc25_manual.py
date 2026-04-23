from __future__ import annotations

import argparse
import json
import re
from dataclasses import dataclass
from datetime import date
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont


MANUAL_STRUCTURE = {
    "title": "BC - Business Central funktionalitet",
    "subtitle": "BC25 slutbrugermanual og opslagsvaerk",
    "domains": [
        {
            "name": "Generelt",
            "intro": "Funktioner i denne sektion bruges pa tvaers af afdelinger og er typisk relevante som daglig opslaasning eller stamdatahjaelp.",
            "entries": [
                {
                    "title": "Rollecentre og Scanpan-genveje",
                    "purpose": "Samler Scanpan-specifikke genveje, statuskort og hurtig adgang til de vigtigste funktioner i rollecentrene.",
                    "location": "Rollecentre som Ordrebehandler, Administrator, Produktion og Lager.",
                    "when": "Bruges som startside for det daglige arbejde, naer brugeren skal have overblik eller ga direkte til en funktion.",
                    "steps": [
                        "Aabn dit rollecenter i Business Central.",
                        "Brug Scanpan-felterne og klynger til at aabne relevante lister, rapporter og overvagningsfelter.",
                        "Klik pa et statusfelt for at gaa direkte til det underliggende dataudsnit, hvis der er drilldown.",
                    ],
                    "important": [
                        "Rollecentrene er opbygget efter arbejdsomraade og skal derfor bruges som primaer indgang til de fleste Scanpan-funktioner.",
                        "Knapper og cue-felter viser kun det, der er relevant for den konkrete rolle.",
                    ],
                    "related": ["Scanpan Process status", "Customers Over Credit Limit", "Sales Comparison & Realized Sales"],
                    "objects": [
                        "pageextension 50020 OrderProcessorRoleCenter",
                        "pageextension 50000 AdminRoleCenterExtSC",
                        "pageextension 50028 ProdPlannerRoleCenterExtSC",
                        "pageextension 50127 WhseBasicRoleCenter",
                        "pageextension 50026 WhseWorkerWMSRoleCenterExtSC",
                    ],
                },
                {
                    "title": "Mandatory Field Setup",
                    "purpose": "Sikrer at paakraevede felter er udfyldt for kunder, leverandoerer og varer, inklusive simple betingede regler.",
                    "location": "Soeg efter `Mandatory Field Setup`.",
                    "when": "Bruges naer virksomheden vil styre datakvalitet og tvinge udfyldelse af bestemte felter i udvalgte forretningsomraader.",
                    "steps": [
                        "Vaelg forretningsomraade i feltet `Business Area`.",
                        "Opret eller rediger de felter, som skal vaere paakraevede.",
                        "Angiv eventuelt conditional-felter og betingelser, hvis kravet kun gaelder i bestemte scenarier.",
                    ],
                    "important": [
                        "Opsaetningen bruges som styring af datakvalitet og boer vedligeholdes centralt.",
                        "Conditional-felter giver mulighed for forskellige krav for forskellige vare- eller kundetyper.",
                    ],
                    "related": ["Fieldpage", "Customer Card", "Item Card"],
                    "objects": ["page 50003 MandatoryFieldSetup", "table 50018 MandatoryFieldSetup2"],
                },
                {
                    "title": "Attachment Overview",
                    "purpose": "Samler vedhaeftede filer for bogfoerte salgsdokumenter, inklusive filer fra baade bogfoerte dokumenter og tilhoerende salgsordrelinjer.",
                    "location": "Soeg efter `Attachment Overview`.",
                    "when": "Bruges naer man hurtigt skal finde alle filer knyttet til en bogfoert faktura eller kreditnota.",
                    "steps": [
                        "Vaelg dokumenttype i `Posted Doc Type`.",
                        "Indtast dokumentnummer i `Posted Doc No.` eller brug `Load Attachments`.",
                        "Gennemgaa resultatlisten og brug `Open Selected` for at aabne eller hente en vedhaeftning.",
                    ],
                    "important": [
                        "Listen er midlertidig og bygger sig op ud fra det dokument, du indtaster.",
                        "Funktionen samler baade header-, linje- og relaterede salgsordrevedhaeftninger i samme overblik.",
                    ],
                    "related": ["Posted Sales Invoice", "Posted Sales Invoices"],
                    "objects": ["page 50095 Attachment Overview"],
                },
                {
                    "title": "Adresseoversigt",
                    "purpose": "Giver en samlet oversigt over kunder, leveringsadresser og leverandoerer med kontakt- og adresseoplysninger.",
                    "location": "Soeg efter `All Customer Addresses` eller aabn siden `Addresses`.",
                    "when": "Bruges ved opslag, kontrol og eksport af adresse- og kontaktdata.",
                    "steps": [
                        "Aabn oversigten.",
                        "Brug filtre til at afgraense pa adresse-type, land, kunde eller leverandoer.",
                        "Eksporter om noedvendigt listen til Excel for videre bearbejdning.",
                    ],
                    "important": [
                        "Oversigten er velegnet til kontrol og masseopslag, ikke til daglig redigering af stamdata.",
                    ],
                    "related": ["Debitor Addresses", "Customer Card", "Vendor Card"],
                    "objects": ["page 50005 Addresses", "query 50004 AddressesCustomer", "query 50005 AddressesVendor"],
                },
                {
                    "title": "Create Multiple Item Barcode",
                    "purpose": "Opretter flere stregkoder for en vare og kan samtidig opdatere GTIN pa alle varer.",
                    "location": "Soeg efter `Create Multiple Item Barcode`.",
                    "when": "Bruges naer der skal oprettes EAN- eller UPC-koder for en vare og dens enheder.",
                    "steps": [
                        "Vaelg varen i `Item to assign barcode`.",
                        "Marker de enheder der skal have EAN eller UPC i listen.",
                        "Koer `Create Barcodes` for at oprette de valgte koder.",
                    ],
                    "important": [
                        "Funktionen rydder den midlertidige arbejdsliste efter oprettelse.",
                        "Knappen `Update all Items GTIN` opdaterer GTIN bredt og boer bruges med omtanke.",
                    ],
                    "related": ["Item Reference Entries", "Item Card"],
                    "objects": ["page 50000 CreateMultipleBarcodes"],
                },
            ],
        },
        {
            "name": "Salg",
            "intro": "Sektionsindholdet stoetter salgsarbejde, kundestyring, ordreopfoelgning og overblik over linjer, priser og kreditforhold.",
            "entries": [
                {
                    "title": "Sales Line",
                    "purpose": "Viser alle salgslinjer i en samlet oversigt med filtre, drilldown og mulighed for videre analyse.",
                    "location": "Soeg efter `Sales Line` eller aabn funktionen fra rollecenteret.",
                    "when": "Bruges naer man vil analysere aabne salgslinjer, afsendelser, restmængder eller kundespecifikke linjer.",
                    "steps": [
                        "Aabn siden og brug filtrene for kunde, saelger, land, enhed eller status.",
                        "Aktiver de relevante toggles for `Outstanding Quantity`, `Shipped Not Invoiced` eller overskrifter.",
                        "Brug drilldown pa dokument, kunde eller vare for at gaa videre til det underliggende opslag.",
                    ],
                    "important": [
                        "Siden er lavet som analyseoversigt og egner sig godt til eksport og opfoelgning.",
                        "Linjevisningen viser blandt andet toldnummer, produktlinje, enhed og beloeb i samme liste.",
                    ],
                    "related": ["Sales Order", "Sales Orders", "Sales Commission"],
                    "objects": ["page 50020 SalesLine", "table 50009 SalesLineTMP"],
                },
                {
                    "title": "Salgsordre",
                    "purpose": "Udvider standard-salgsordren med Scanpan-felter, advarsler, faktabokse og en handling til kopiering fra Sell-to til Ship-to/Bill-to.",
                    "location": "Standardsiden `Sales Order` med Scanpan-udvidelser.",
                    "when": "Bruges ved registrering og vedligeholdelse af salgsordrer, isaar hvor leverings- og faktureringsoplysninger skal kopieres eller kontrolleres.",
                    "steps": [
                        "Aabn salgsordren som normalt.",
                        "Brug handlingen `Copy Sell-to to Ship-to/Bill-to`, hvis leverings- og faktureringsadresser skal overstyres med oplysninger fra sell-to felterne.",
                        "Gennemgaa faktabokse og advarsler inden ordren frigives eller behandles videre.",
                    ],
                    "important": [
                        "Kopieringshandlingen overskriver eksisterende shipping- og billing-oplysninger efter bekraeftelse.",
                        "Ordren viser ogsaa systemoplysninger om oprettelse og aendring samt advarsler fra Scanpan-logik.",
                    ],
                    "related": ["Sales Line", "Posted Sales Invoice", "Customers Over Credit Limit"],
                    "objects": ["pageextension 50012 SalesOrder"],
                },
                {
                    "title": "Customers Over Credit Limit",
                    "purpose": "Viser kunder der overskrider kredit- eller selvrisikograense og kan skifte mellem detalje- og totalvisning.",
                    "location": "Cue/part pa rollecentre som Ordrebehandler og Administrator.",
                    "when": "Bruges til daglig opfoelgning pa kreditrisiko foer ordrebehandling eller frigivelse.",
                    "steps": [
                        "Aabn cue-feltet fra rollecenteret.",
                        "Skift mellem detaljer og summer med `Toggle View`.",
                        "Brug resultatet som beslutningsstoette foer videre salgsbehandling.",
                    ],
                    "important": [
                        "Visningen viser baade almindelig kreditgraense og self-insured limit.",
                        "Funktionen er lavet til overvagning og ikke som erstatning for kreditpolitik.",
                    ],
                    "related": ["Sales Order", "Scanpan Setup warning settings"],
                    "objects": ["page 50047 CustomersOverCreditLimit"],
                },
                {
                    "title": "Sales Comparison & Realized Sales",
                    "purpose": "Giver salgsoversigt og maalinger direkte i rollecenteret som et hurtigt ledelses- og opfoelgningsbillede.",
                    "location": "CardPart i rollecentre.",
                    "when": "Bruges naer salg eller ledelse vil have hurtigt statusbillede uden at aabne flere analyser.",
                    "steps": [
                        "Aabn rollecenteret.",
                        "Se de viste maalinger og klik videre, hvor der er drilldown.",
                        "Brug informationen til hurtig opfoelgning eller videre analyse i andre lister.",
                    ],
                    "important": [
                        "Kortet er et overblik og ikke den detaljerede analyse i sig selv.",
                    ],
                    "related": ["Sales Line", "Rollecentre og Scanpan-genveje"],
                    "objects": ["page 50060 SalesCompareAndRealized"],
                },
                {
                    "title": "Inter Company Tracking",
                    "purpose": "Samler sporingsinformation for intercompany-flow mellem blandt andet Danmark og Norge.",
                    "location": "Soeg efter `Inter Company Information Tracking` eller aabn fra rollecenterets salgsgenveje.",
                    "when": "Bruges ved opfoelgning pa dokumentflow mellem selskaber og ved fejlsoegning i intercompany-processer.",
                    "steps": [
                        "Aabn listen.",
                        "Filtrer pa relevante dokumenter eller forloeb.",
                        "Brug listen til at identificere manglende eller forsinkede trin i intercompany-processen.",
                    ],
                    "important": [
                        "Funktionen er et kontrol- og sporingsvaerktoej for intercompany, ikke en erstatning for bogfoeringskontrol.",
                    ],
                    "related": ["Scanpan Process status"],
                    "objects": ["page 50022 IICTracking_BC"],
                },
                {
                    "title": "Sales Commission",
                    "purpose": "Viser kommissionsrelaterede oplysninger pr. vare og understoetter opfoelgning pa saelgerperformance.",
                    "location": "Soeg efter `Sales Commission`.",
                    "when": "Bruges ved analyse og kontrol af kommissionsgrundlag.",
                    "steps": [
                        "Aabn listen.",
                        "Filtrer pa relevante varer, perioder eller saelgere.",
                        "Brug resultatet som grundlag for opfoelgning eller videre rapportering.",
                    ],
                    "important": [
                        "Funktionen er beregnet som analysevisning og boer ses sammen med almindelig salgsrapportering.",
                    ],
                    "related": ["Sales Line", "Sales Comparison & Realized Sales"],
                    "objects": ["page 50035 Sales Commission"],
                },
            ],
        },
        {
            "name": "Fakturering og Claims",
            "intro": "Her ligger funktioner, der stoetter bogfoerte dokumenter, toldoplysninger og den saerskilte claims-loesning.",
            "entries": [
                {
                    "title": "Customs Declaration",
                    "purpose": "Udtraekker varedata til told- og eksportrelateret arbejde med filter, oversat beskrivelse og mulighed for rapportkoersel.",
                    "location": "Soeg efter `Customs Declaration` eller aabn funktionen fra rollecenterets salgsrapporter.",
                    "when": "Bruges ved arbejde med told, eksportdokumentation og varedata til myndigheds- eller kundeformaal.",
                    "steps": [
                        "Aabn siden og vaelg eventuelt varebeskrivelsens sprog.",
                        "Brug filtrene for bogfoeringsgruppe, brugsgruppe, ABCD-kategori, varenummer eller kvalitet.",
                        "Koer rapporten eller brug listevisningen som grundlag for eksport.",
                    ],
                    "important": [
                        "Beskrivelsen hentes paa det valgte oversaettelsessprog, men falder tilbage til varens egen beskrivelse, hvis oversaettelse mangler.",
                        "Listen viser blandt andet maal, vaegt, cubage, toldnummer og oprindelsesland.",
                    ],
                    "related": ["Attachment Overview", "Posted Sales Invoice"],
                    "objects": ["page 50199 Customs Declaration List", "report 50002 Customs Declaration"],
                },
                {
                    "title": "Posted Sales Invoice og Posted Sales Invoices",
                    "purpose": "Udvider bogfoerte salgsfakturaer med ekstra soegefelter, web-oplysninger og handling til overfoersel af nye fakturaer til LTS.",
                    "location": "Standardsiderne `Posted Sales Invoice` og `Posted Sales Invoices` med Scanpan-udvidelser.",
                    "when": "Bruges ved opslag paa bogfoerte fakturaer, webordrer og opfoelgning pa overfoersel til eksterne systemer.",
                    "steps": [
                        "Aabn den bogfoerte faktura eller listen over bogfoerte fakturaer.",
                        "Brug de ekstra felter som `Order No.`, `External Document No.` og `Drop Shipment` til soegning og kontrol.",
                        "Koer `Send all new Invoices to LTS`, hvis nye fakturaer skal sendes videre til LTS.",
                    ],
                    "important": [
                        "Visningen indeholder ogsaa webordre-oplysninger som e-mail, telefon og Payment ID.",
                        "Listevisningen har en dedikeret visning til fakturaer, der endnu ikke er sendt til toll-system.",
                    ],
                    "related": ["Attachment Overview", "Customs Declaration"],
                    "objects": ["pageextension 50056 PostedSalesInvoice", "pageextension 50054 PostedSalesInvoices"],
                },
                {
                    "title": "Claims Admin og Claims Setup",
                    "purpose": "Samler opsaetning af claims-loesningen, inklusive tilladte returarsager, oversaettelser, aarskoder og regler for vareopslag.",
                    "location": "Soeg efter `Claims Admin` eller `Claims Setup`.",
                    "when": "Bruges af administratorer eller procesansvarlige, der vedligeholder claims-processen og dens regler.",
                    "steps": [
                        "Aabn `Claims Admin` for det samlede overblik.",
                        "Vedligehold grundopsaetning som `ClaimsEnabled`, filtrering af varegrupper og fallback-sprog.",
                        "Vedligehold underliggende lister for returarsager, oversaettelser, produktbrug og aarskoder fra de indbyggede dele eller genvejshandlinger.",
                    ],
                    "important": [
                        "Claims anvender Business Central-returarsager som grundlag for fejl- og aarsagskoder.",
                        "Engelsk boer vedligeholdes som oversaettelse for alle returarsager, der bruges i claims.",
                    ],
                    "related": ["Claim Year Codes", "Claim Product Usage Reasons"],
                    "objects": ["page 50289 ClaimsAdmin", "page 50270 ClaimsSetupCard"],
                },
                {
                    "title": "Invoice TariffCode",
                    "purpose": "Rapport til visning af varekode- og toldrelaterede oplysninger pa fakturagrundlag.",
                    "location": "Rapporten `Invoice TariffCode` / `Faktura Varekoder`.",
                    "when": "Bruges naer der skal dokumenteres told- eller varekodeoplysninger pa fakturalinjer.",
                    "steps": [
                        "Aabn rapporten fra soegning eller rollecenter, hvis den er gjort tilgaengelig der.",
                        "Angiv de relevante filtre for faktura eller periode.",
                        "Koer rapporten og gennemgaa resultatet i layoutet.",
                    ],
                    "important": [
                        "Rapporten er dokumentorienteret og boer bruges som udtraek, ikke som redigeringsflade.",
                    ],
                    "related": ["Customs Declaration", "Posted Sales Invoice"],
                    "objects": ["report 50009 Faktura Varekoder"],
                },
            ],
        },
        {
            "name": "Lager og Logistik",
            "intro": "Domænet samler funktioner til lagerstyring, pluk, forsendelser, distributionsregler og vareopslag i lageret.",
            "entries": [
                {
                    "title": "Warehouse Shipment List",
                    "purpose": "Udvider lagerleverancelisten med ekstra transport- og forsendelsesoplysninger samt genvej til lagerbalance for pluk.",
                    "location": "Standardsiden `Warehouse Shipment List` med Scanpan-udvidelser.",
                    "when": "Bruges ved daglig forsendelsesopfoelgning og kontrol af lagerleverancer.",
                    "steps": [
                        "Aabn lagerleverancelisten.",
                        "Brug de ekstra kolonner som beskrivelse, picknummer, shipping agent og servicekode til at finde de rigtige leverancer.",
                        "Koer `Warehouse Pick Balance List` for at kontrollere balance mellem pluk og beholdning.",
                    ],
                    "important": [
                        "Listen aabner med seneste nummer foerst for at forbedre daglig opfoelgning.",
                        "Kolonnerne er tilfoejet for at samle transportinformation pa samme liste uden ekstra opslag.",
                    ],
                    "related": ["Warehouse Pick Bin Balance", "Warehouse Shipment"],
                    "objects": ["pageextension 50046 WarehouseShipmentList"],
                },
                {
                    "title": "Warehouse Pick Bin Balance",
                    "purpose": "Viser forskellen mellem plukket mængde og faktisk lokationsbeholdning pr. bin.",
                    "location": "Soeg efter `Warehouse Pick Bin Balance` eller aabn fra lagerleverancelisten.",
                    "when": "Bruges naer lageret skal finde ubalancer mellem pluk og fysisk beholdning.",
                    "steps": [
                        "Aabn listen.",
                        "Angiv eventuelt `Date Filter` for det relevante udsnit.",
                        "Gennemgaa linjer med negativ balance, som markeres visuelt i listen.",
                    ],
                    "important": [
                        "Negativ `Bin Quantity Balance` er et klart kontrolsignal og boer undersoeges.",
                    ],
                    "related": ["Warehouse Shipment List", "Warehouse Picks"],
                    "objects": ["page 50019 WMSPickBinBalance"],
                },
                {
                    "title": "Bin Contents Scanpan-filter",
                    "purpose": "Tilfoejer lagerfiltre til bin-indhold, sa det er lettere at finde varer til Auning, kun bins med indhold og linjer med transfer orders.",
                    "location": "Standardsiden `Bin Contents` med Scanpan-udvidelser.",
                    "when": "Bruges ved lageropfoelgning, transfer-forberedelse og hurtig filtrering af lagerindhold.",
                    "steps": [
                        "Aabn `Bin Contents`.",
                        "Brug `Set AUNING Filters`, `Show only with Contents` og `Show only Transfer Orders` efter behov.",
                        "Gennemgaa resultatet med de ekstra Scanpan-kolonner som produktlinje og transfer order.",
                    ],
                    "important": [
                        "Auning-filteret saetter baade lagersted og inventory posting group.",
                        "Filtrene kan kombineres for at afgraense lagerudsnittet hurtigt.",
                    ],
                    "related": ["Transfer Orders", "Warehouse Shipment List"],
                    "objects": ["pageextension 50053 BinContentsExtSC"],
                },
                {
                    "title": "Shipping Agent Distribution",
                    "purpose": "Vedligeholder fordelingsregler for transportoerer, land og intervalomraader.",
                    "location": "Soeg efter `Shipping Agent Distribution`.",
                    "when": "Bruges naer distributionsregler for fragtagenter eller postnummere skal opdateres.",
                    "steps": [
                        "Aabn listen over distributionsregler.",
                        "Opret eller rediger poster med transportoer, land og interval.",
                        "Brug `New Distribution Entry` som genvej til oprettelse.",
                    ],
                    "important": [
                        "Reglerne boer vedligeholdes centralt, da de kan paavirke forsendelsesflow og agentvalg.",
                    ],
                    "related": ["Shipping Agents", "Shipping Agent Services"],
                    "objects": ["page 50058 ShipAgentDistribution"],
                },
            ],
        },
        {
            "name": "Produktion og planlaegning",
            "intro": "Omraaderne her bruges af planlaegning, produktion og controlling til overblik, prioritering og detaljeanalyse.",
            "entries": [
                {
                    "title": "Production Controlling Lines",
                    "purpose": "Samler transferlinjer og produktionslinjer i en kontrolleret analysevisning for produktionen.",
                    "location": "Soeg efter `Production Controlling Lines`.",
                    "when": "Bruges ved daglig controlling og prioritering af aabne eller udfordrede produktionsforloeb.",
                    "steps": [
                        "Aabn siden og vaelg om transferlinjer, firm planned og released produktionslinjer skal med.",
                        "Angiv eventuelt varslingsafstand via `Dateformula Ending Date`.",
                        "Koer `Fetch Lines` for at opdatere visningen.",
                    ],
                    "important": [
                        "Listen er en beregnet analysevisning og skal opdateres med `Fetch Lines` for at vaere aktuel.",
                        "Linjer markeres visuelt ved kritiske datoer eller negative resterende produktionsenheder.",
                    ],
                    "related": ["Production Controlling Dashboard", "Production Controlling PanPlan"],
                    "objects": ["page 50013 ProductionControlling"],
                },
                {
                    "title": "Production Controlling Dashboard",
                    "purpose": "Viser grafer og kapacitetsbillede for produktionen og giver drilldown fra datapunkter til konkrete produktionsordrer.",
                    "location": "Soeg efter `Production Controlling Dashboard`.",
                    "when": "Bruges naer produktionen skal have visuelt overblik over presser, maskiner og belastning.",
                    "steps": [
                        "Aabn dashboardet.",
                        "Gennemgaa de viste grafer for Foundry og Processing.",
                        "Klik pa et datapunkt for at slaa de bagvedliggende produktionsordrer op.",
                    ],
                    "important": [
                        "Dashboardet er et visuelt indgangspunkt og fungerer bedst sammen med de detaljerede controlling-lister.",
                    ],
                    "related": ["Production Controlling Lines", "Production Controlling Routing List"],
                    "objects": ["page 50044 ProdControllingDashboard"],
                },
                {
                    "title": "Production Controlling PanPlan",
                    "purpose": "Bygger et controllingudsnit baseret pa produktionsordrer, styklister, lager og indkoeb til PanPlan-opfoelgning.",
                    "location": "Soeg efter `SCANPAN Production Controlling PanPlan`.",
                    "when": "Bruges ved dybere produktionsanalyse, isaar hvor styklister og komponentniveauer skal sammenholdes med ordrer og lager.",
                    "steps": [
                        "Aabn siden.",
                        "Laes linjerne som et genereret overblik over produktionsordre, komponent og lagermaengde.",
                        "Brug sortering og filtre til at fokusere pa de kritiske forloeb.",
                    ],
                    "important": [
                        "Datagrundlaget genereres ved aabning og bygger pa aktiv produktion og komponenter.",
                        "Visningen er analyseorienteret og ikke lavet til direkte transaktionsbehandling.",
                    ],
                    "related": ["Production Controlling Lines", "Prod. Controlling Recursive BOM List"],
                    "objects": ["page 50030 ProdControllingPanPlan"],
                },
                {
                    "title": "Recursive BOM og BOM-vaerktoejer",
                    "purpose": "Stoetter analyse og vedligeholdelse af styklister, omkostningsandele og justering af BOM-linjer.",
                    "location": "Soeg efter `Prod. Controlling Recursive BOM List`, `Recursive BOM Cost Shares` eller `Adjust BOM Lines`.",
                    "when": "Bruges naer produktion eller controlling skal analysere eller justere styklisteopbygning og omkostninger.",
                    "steps": [
                        "Aabn den relevante BOM-funktion afhængigt af opgaven.",
                        "Gennemgaa stykliste- eller omkostningsvisningen.",
                        "Koer justering eller analyse efter behov.",
                    ],
                    "important": [
                        "Disse funktioner bruges typisk af avancerede brugere i produktion eller controlling.",
                    ],
                    "related": ["Production Controlling Lines", "Production Controlling PanPlan"],
                    "objects": [
                        "page 50061 ProdControllingRecursiveBOM",
                        "page 50069 BOMCostShares",
                        "page 50054 AdjustBoMlines",
                    ],
                },
            ],
        },
        {
            "name": "Kampagner",
            "intro": "Kampagnefunktionerne samler genveje, kommentarer og salgsopfoelgning omkring kampagnearbejde.",
            "entries": [
                {
                    "title": "Kampagner og kampagnesalg",
                    "purpose": "Understoetter kampagnearbejde med genveje fra rollecenter, kommentarer og visning af kampagnesalg.",
                    "location": "Kampagnegenveje pa rollecenteret samt siderne `Campaign List`, `Campaign Card`, `Campaign Comments` og `Sales Campaign Sales`.",
                    "when": "Bruges af salg og marketing til oprettelse, vedligeholdelse og opfoelgning pa kampagner.",
                    "steps": [
                        "Aabn kampagnelisten eller kampagnekortet fra rollecenterets genveje.",
                        "Vedligehold kommentarer i kampagnens kommentarfelt eller faktaboks.",
                        "Brug kampagnesalgsvisningen til opfoelgning pa resultater.",
                    ],
                    "important": [
                        "Rollecenteret indeholder hurtig adgang til kampagnelister, segmenter og mailgrupper.",
                    ],
                    "related": ["Campaign List", "Campaign Comments", "Sales Campaign Sales"],
                    "objects": [
                        "pageextension 50016 CampaignList",
                        "pageextension 50082 CampaignCard",
                        "page 50067 CampaignCommentsFactBox",
                        "page 50037 CampaignSales",
                    ],
                }
            ],
        },
        {
            "name": "Administration og opsaetning",
            "intro": "Denne sektion samler overvaagning, advarselsindstillinger og administrative opsaetninger, som stoetter den daglige drift.",
            "entries": [
                {
                    "title": "Scanpan Process status",
                    "purpose": "Viser job queue-status, intercompany-bufferstatus og pending sync entries direkte i rollecenteret.",
                    "location": "CardPart pa udvalgte rollecentre.",
                    "when": "Bruges af administratorer og superbrugere til daglig overvagning af integrationer og batchkørsel.",
                    "steps": [
                        "Aabn rollecenteret.",
                        "Gennemgaa felterne for Job Queue, InterCompany og MasterData.",
                        "Klik pa et felt for at gaa direkte til de poster, der skal undersoeges.",
                    ],
                    "important": [
                        "Drilldown virker kun korrekt i det relevante selskab og giver fejlbesked, hvis brugeren staar i forkert company.",
                    ],
                    "related": ["Job Queue Entries", "Inter Company Tracking"],
                    "objects": ["page 50041 ScanpanCardPart"],
                },
                {
                    "title": "SCANPAN Setup warning settings",
                    "purpose": "Styrer om brugerne skal se advarsler ved self-insured og kreditmaksimum pa salgsordrer.",
                    "location": "Standardsiden `SCANPAN Setup` med Scanpan-udvidelse.",
                    "when": "Bruges ved fastsaettelse af, hvordan kreditrelaterede advarsler skal opfoere sig i salgsprocessen.",
                    "steps": [
                        "Aabn `SCANPAN Setup`.",
                        "Gaa til gruppen `Warning Settings`.",
                        "Vedligehold felterne for `Show SelfInsured Warning` og `Show CreditMax Warning`.",
                    ],
                    "important": [
                        "Opsaetningen paavirker brugeroplevelsen pa salgsordrer og boer aendres kontrolleret.",
                    ],
                    "related": ["Sales Order", "Customers Over Credit Limit"],
                    "objects": ["pageextension 50120 ScanpanSetup"],
                },
            ],
        },
        {
            "name": "Integrationer og dataudtraek",
            "intro": "Funktionerne i denne sektion bruges ved systemudveksling, webservices og avancerede dataudtraek. De er typisk relevante for superbrugere og integrationsansvarlige.",
            "entries": [
                {
                    "title": "WebService Sales Pricelist Source Data",
                    "purpose": "Udtraekker vare-, oversaettelses-, stregkode- og prisdata til ekstern bearbejdning.",
                    "location": "Soeg efter `WebSevice Sales Pricelist Source Data`.",
                    "when": "Bruges naer eksterne loesninger eller avancerede prisprocesser har brug for et standardiseret dataudtraek fra BC.",
                    "steps": [
                        "Aabn listen.",
                        "Filtrer pa relevante varer om noedvendigt.",
                        "Eksporter data til den videre eksterne proces.",
                    ],
                    "important": [
                        "Listen viser baade oversaettelser, barcode, enheder, indkøbspris og valuta.",
                        "Funktionen er et dataudtraek og ikke en redigeringsside.",
                    ],
                    "related": ["Sales Price Lists", "Customs Declaration"],
                    "objects": ["page 50028 WebServiceSalesPriceListSource"],
                },
                {
                    "title": "DSV API og transportdata",
                    "purpose": "Understoetter transport- og containerrelaterede dataudtraek og teknisk kontrol omkring DSV-forloeb.",
                    "location": "Soeg efter `DSV API`.",
                    "when": "Bruges ved integration, kontrol eller fejlfinding omkring DSV-transportdata.",
                    "steps": [
                        "Aabn DSV API-siden.",
                        "Filtrer eller gennemgaa de relevante poster.",
                        "Brug resultatet i forbindelse med integration eller kontrol.",
                    ],
                    "important": [
                        "Funktionen er mest relevant for superbrugere og integrationsansvarlige.",
                    ],
                    "related": ["Shipping Agent Distribution", "Warehouse Shipment List"],
                    "objects": ["page 50007 DSVAPI"],
                },
            ],
        },
    ],
}

DOMAIN_COLORS = {
    "Generelt": ("#1F2937", "#9CA3AF"),
    "Salg": ("#7C2D12", "#F59E0B"),
    "Fakturering og Claims": ("#1E3A8A", "#60A5FA"),
    "Lager og Logistik": ("#14532D", "#86EFAC"),
    "Produktion og planlaegning": ("#4C1D95", "#C4B5FD"),
    "Kampagner": ("#9A3412", "#FDBA74"),
    "Administration og opsaetning": ("#0F766E", "#99F6E4"),
    "Integrationer og dataudtraek": ("#374151", "#D1D5DB"),
}


@dataclass(frozen=True)
class ObjectInfo:
    kind: str
    object_id: int
    name: str
    extends: str
    caption: str
    usage_category: str
    page_type: str
    file_path: str


def humanize(text: str) -> str:
    replacements = (
        ("Aae", "Åe"),
        ("Ooe", "Øe"),
        ("Aa", "Å"),
        ("aa", "å"),
        ("Ae", "Æ"),
        ("ae", "æ"),
        ("Oe", "Ø"),
        ("oe", "ø"),
    )
    result = text
    for source, target in replacements:
        result = result.replace(source, target)
    return result


def get_font(size: int) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    font_candidates = [
        Path(r"C:\Windows\Fonts\Aptos.ttf"),
        Path(r"C:\Windows\Fonts\arial.ttf"),
        Path(r"C:\Windows\Fonts\segoeui.ttf"),
    ]
    for candidate in font_candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


def hex_to_rgb(value: str) -> tuple[int, int, int]:
    value = value.lstrip("#")
    return tuple(int(value[index:index + 2], 16) for index in (0, 2, 4))


def create_domain_banner(assets_dir: Path, domain_name: str, intro: str) -> Path:
    assets_dir.mkdir(parents=True, exist_ok=True)
    output_path = assets_dir / f"{domain_name.lower().replace(' ', '_')}.png"

    primary_hex, accent_hex = DOMAIN_COLORS.get(domain_name, ("#1F2937", "#9CA3AF"))
    primary = hex_to_rgb(primary_hex)
    accent = hex_to_rgb(accent_hex)

    width, height = 1600, 420
    image = Image.new("RGB", (width, height), primary)
    draw = ImageDraw.Draw(image)

    for x in range(width):
        ratio = x / max(width - 1, 1)
        blended = tuple(int(primary[index] * (1 - ratio) + accent[index] * ratio) for index in range(3))
        draw.line((x, 0, x, height), fill=blended)

    draw.ellipse((width - 420, -80, width + 120, 380), fill=(255, 255, 255, 26))
    draw.ellipse((width - 260, 120, width + 140, 520), fill=(255, 255, 255, 18))
    draw.rounded_rectangle((70, 60, 260, 118), radius=24, fill=(255, 255, 255))
    draw.text((100, 76), "BC25", fill=primary, font=get_font(30))
    draw.text((70, 155), humanize(domain_name), fill="white", font=get_font(54))

    intro_font = get_font(24)
    lines: list[str] = []
    words = humanize(intro).split()
    current = ""
    for word in words:
        proposal = f"{current} {word}".strip()
        if draw.textlength(proposal, font=intro_font) <= 920:
            current = proposal
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)

    y = 245
    for line in lines[:3]:
        draw.text((72, y), line, fill=(245, 245, 245), font=intro_font)
        y += 34

    image.save(output_path)
    return output_path


def create_cover_image(assets_dir: Path) -> Path:
    assets_dir.mkdir(parents=True, exist_ok=True)
    custom_cover_path = assets_dir / "manual_cover_custom.png"
    if custom_cover_path.exists():
        return custom_cover_path
    output_path = assets_dir / "manual_cover.png"

    width, height = 1800, 980
    image = Image.new("RGB", (width, height), hex_to_rgb("#E9E2D8"))
    draw = ImageDraw.Draw(image)

    top_color = hex_to_rgb("#2B313B")
    mid_color = hex_to_rgb("#4F5B66")
    bottom_color = hex_to_rgb("#C46A3B")
    for y in range(height):
        ratio = y / max(height - 1, 1)
        if ratio < 0.65:
            inner_ratio = ratio / 0.65
            blended = tuple(int(top_color[index] * (1 - inner_ratio) + mid_color[index] * inner_ratio) for index in range(3))
        else:
            inner_ratio = (ratio - 0.65) / 0.35
            blended = tuple(int(mid_color[index] * (1 - inner_ratio) + bottom_color[index] * inner_ratio) for index in range(3))
        draw.line((0, y, width, y), fill=blended)

    draw.ellipse((1180, -130, 1760, 450), fill=(255, 255, 255, 36))
    draw.ellipse((980, 160, 1560, 740), fill=(255, 255, 255, 22))

    panel_fill = (246, 243, 238)
    panel_outline = (212, 199, 183)
    dashboard = (950, 150, 1660, 620)
    draw.rounded_rectangle(dashboard, radius=36, fill=panel_fill, outline=panel_outline, width=2)
    draw.rounded_rectangle((1010, 215, 1320, 410), radius=24, fill=(255, 255, 255), outline=(220, 220, 220), width=2)
    draw.rounded_rectangle((1350, 215, 1600, 330), radius=24, fill=(255, 255, 255), outline=(220, 220, 220), width=2)
    draw.rounded_rectangle((1350, 355, 1600, 555), radius=24, fill=(255, 255, 255), outline=(220, 220, 220), width=2)

    draw.text((1040, 245), "BC25", fill="#2B313B", font=get_font(30))
    draw.text((1040, 292), "Sales & Operations", fill="#2B313B", font=get_font(36))

    chart_points = [(1040, 380), (1105, 330), (1180, 350), (1250, 275)]
    draw.line(chart_points, fill="#C46A3B", width=6)
    for x, y in chart_points:
        draw.ellipse((x - 8, y - 8, x + 8, y + 8), fill="#C46A3B")
    draw.text((1378, 242), "Warehouse", fill="#1F2937", font=get_font(24))
    for idx, bar_height in enumerate([40, 72, 58]):
        x0 = 1392 + idx * 56
        draw.rounded_rectangle((x0, 308 - bar_height, x0 + 28, 308), radius=8, fill="#5AA6D6")
    draw.text((1378, 382), "Planning", fill="#1F2937", font=get_font(24))
    for idx, value in enumerate([54, 84, 62, 98]):
        x0 = 1390 + idx * 46
        draw.rounded_rectangle((x0, 520 - value, x0 + 24, 520), radius=8, fill="#7F8A3B")

    for shelf_y in (640, 720, 800):
        draw.line((960, shelf_y, 1700, shelf_y), fill=(80, 66, 56), width=6)
    for shelf_x in (1020, 1180, 1340, 1500, 1660):
        draw.line((shelf_x, 610, shelf_x, 835), fill=(80, 66, 56), width=6)
    for x0, y0, x1, y1, color in (
        (990, 664, 1140, 714, "#B65E35"),
        (1160, 664, 1310, 714, "#7F8A3B"),
        (1330, 664, 1480, 714, "#D49A42"),
        (1500, 664, 1650, 714, "#4E7FA5"),
        (1030, 744, 1220, 794, "#C46A3B"),
        (1260, 744, 1450, 794, "#65735B"),
        (1490, 744, 1680, 794, "#2F4B6C"),
    ):
        draw.rounded_rectangle((x0, y0, x1, y1), radius=12, fill=color)

    draw.line((970, 578, 1080, 642), fill="#D9B873", width=5)
    draw.line((1078, 640, 1220, 610), fill="#D9B873", width=5)
    draw.line((1218, 610, 1370, 660), fill="#D9B873", width=5)
    draw.line((1368, 658, 1528, 620), fill="#D9B873", width=5)
    draw.line((1526, 620, 1660, 690), fill="#D9B873", width=5)
    for point in ((970, 578), (1080, 642), (1220, 610), (1370, 660), (1528, 620), (1660, 690)):
        x, y = point
        draw.ellipse((x - 7, y - 7, x + 7, y + 7), fill="#F6E7B5")

    body_color = "#21262D"
    rim_color = "#AAB2BD"
    copper = "#C7794D"
    pan1 = (110, 520, 560, 860)
    draw.ellipse(pan1, fill=body_color, outline=rim_color, width=8)
    draw.ellipse((160, 565, 510, 815), fill="#353C45", outline="#616B77", width=4)
    draw.rectangle((500, 650, 810, 702), fill=body_color)
    draw.rounded_rectangle((780, 640, 1010, 710), radius=28, fill=copper)
    pan2 = (280, 330, 820, 700)
    draw.ellipse(pan2, fill=body_color, outline=rim_color, width=8)
    draw.ellipse((335, 382, 765, 648), fill="#39424D", outline="#66717F", width=4)
    draw.rectangle((760, 470, 1060, 526), fill=body_color)
    draw.rounded_rectangle((1030, 458, 1210, 540), radius=30, fill=copper)
    lid = (410, 225, 690, 355)
    draw.ellipse(lid, fill="#454F5B", outline="#B8C0CA", width=5)
    draw.rounded_rectangle((520, 192, 580, 234), radius=16, fill=copper)

    heading_font = get_font(66)
    sub_font = get_font(30)
    draw.text((86, 90), "SCANPAN BC25", fill="white", font=heading_font)
    draw.text((88, 178), "Slutbrugervejledning med fokus på salg, lager, logistik og produktion", fill=(241, 234, 226), font=sub_font)
    draw.rounded_rectangle((88, 252, 350, 312), radius=18, fill=(245, 240, 233))
    draw.text((118, 268), "Business Central", fill="#2B313B", font=get_font(28))
    draw.rounded_rectangle((370, 252, 515, 312), radius=18, fill=(245, 240, 233))
    draw.text((410, 268), "BC25", fill="#2B313B", font=get_font(28))

    image.save(output_path)
    return output_path


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_callout_table(document: Document, rows: list[tuple[str, str]], fill: str = "F3F4F6") -> None:
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        row = table.add_row().cells
        row[0].text = humanize(label)
        row[1].text = humanize(value)
        set_cell_shading(row[0], fill)
        row[0].paragraphs[0].runs[0].bold = True
        row[0].width = Inches(1.8)
    document.add_paragraph()


OBJECT_PATTERNS = {
    "page": re.compile(r"^\s*page\s+(\d+)\s+\"?([^\"]+?)\"?\s*$", re.MULTILINE),
    "pageextension": re.compile(r"^\s*pageextension\s+(\d+)\s+\"?([^\"]+?)\"?\s+extends\s+\"?([^\"]+?)\"?\s*$", re.MULTILINE),
    "report": re.compile(r"^\s*report\s+(\d+)\s+\"?([^\"]+?)\"?\s*$", re.MULTILINE),
    "reportextension": re.compile(r"^\s*reportextension\s+(\d+)\s+\"?([^\"]+?)\"?\s+extends\s+\"?([^\"]+?)\"?\s*$", re.MULTILINE),
}


def strip_comments(text: str) -> str:
    without_block_comments = re.sub(r"/\*.*?\*/", "", text, flags=re.DOTALL)
    return re.sub(r"//.*", "", without_block_comments)


def get_match(text: str) -> tuple[str, re.Match[str]] | None:
    for kind, pattern in OBJECT_PATTERNS.items():
        match = pattern.search(text)
        if match:
            return kind, match
    return None


def extract_property(text: str, property_name: str) -> str:
    pattern = re.compile(rf"\b{re.escape(property_name)}\s*=\s*'([^']*)'")
    match = pattern.search(text)
    if match:
        return match.group(1).strip()
    pattern = re.compile(rf"\b{re.escape(property_name)}\s*=\s*([A-Za-z0-9_.]+)")
    match = pattern.search(text)
    if match:
        return match.group(1).strip()
    return ""


def collect_inventory(repo_root: Path) -> list[ObjectInfo]:
    src_root = repo_root / "src"
    inventory: list[ObjectInfo] = []

    for path in sorted(src_root.rglob("*.al")):
        raw_text = path.read_text(encoding="utf-8-sig", errors="ignore")
        text = strip_comments(raw_text)
        match_data = get_match(text)
        if not match_data:
            continue

        kind, match = match_data
        groups = match.groups()
        object_id = int(groups[0])
        name = groups[1].strip()
        extends = groups[2].strip() if len(groups) > 2 else ""

        inventory.append(
            ObjectInfo(
                kind=kind,
                object_id=object_id,
                name=name,
                extends=extends,
                caption=extract_property(text, "Caption"),
                usage_category=extract_property(text, "UsageCategory"),
                page_type=extract_property(text, "PageType"),
                file_path=str(path.relative_to(repo_root)),
            )
        )

    return inventory


def configure_styles(document: Document) -> None:
    document.styles["Normal"].font.name = "Aptos"
    document.styles["Normal"].font.size = Pt(10.5)

    for style_name, size in (("Title", 24), ("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)):
        style = document.styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(size)

    if "Manual Label" not in document.styles:
        label_style = document.styles.add_style("Manual Label", WD_STYLE_TYPE.PARAGRAPH)
        label_style.base_style = document.styles["Normal"]
        label_style.font.bold = True
        label_style.font.size = Pt(10.5)


def add_field_code(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")

    result = OxmlElement("w:t")
    result.text = humanize("Hoejreklik og opdater felt for indholdsfortegnelse.")

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(result)
    run._r.append(fld_char_end)


def add_label_paragraph(document: Document, label: str, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.add_run(f"{humanize(label)}: ").bold = True
    paragraph.add_run(humanize(text))


def add_bullets(document: Document, items: Iterable[str]) -> None:
    for item in items:
        document.add_paragraph(humanize(item), style="List Bullet")


def add_numbered(document: Document, items: Iterable[str]) -> None:
    for item in items:
        document.add_paragraph(humanize(item), style="List Number")


def add_cover(document: Document, app_info: dict[str, object], repo_root: Path) -> None:
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.style = "Title"
    title.add_run(humanize(str(MANUAL_STRUCTURE["title"])))

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(humanize(str(MANUAL_STRUCTURE["subtitle"]))).italic = True

    cover_image = create_cover_image(document._manual_assets_dir)
    image_paragraph = document.add_paragraph()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.add_run().add_picture(str(cover_image), width=Inches(6.4))

    intro_box = document.add_table(rows=1, cols=1)
    intro_box.style = "Table Grid"
    intro_box.cell(0, 0).text = humanize(
        "En samlet guide til de vigtigste Scanpan-funktioner i Business Central. Manualen er skrevet til hverdagsbrug og fokuserer paa overblik, arbejdsglaede og hurtig adgang til de funktioner, der goer arbejdet lettere."
    )
    set_cell_shading(intro_box.cell(0, 0), "F7F3EE")
    document.add_paragraph()

    metadata = document.add_table(rows=0, cols=2)
    metadata.style = "Table Grid"
    for label, value in (
        ("Gaelder for", "BC25"),
        ("App", str(app_info.get("name", ""))),
        ("Version", str(app_info.get("version", ""))),
        ("Maalgruppe", "Slutbrugere, superbrugere og procesansvarlige"),
        ("Genereret", date.today().isoformat()),
        ("Kilde", "Aktive objekter i projektbiblioteket under src/"),
    ):
        row = metadata.add_row().cells
        row[0].text = humanize(label)
        row[1].text = value
        row[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    document.add_paragraph(
        humanize(
            "Manualen er genopbygget ud fra det aktuelle BC25-projektbibliotek og er struktureret efter forretningsdomaener, saa den kan bruges som et levende opslagsvaerk for slutbrugere."
        )
    )
    document.add_section(WD_SECTION.NEW_PAGE)


def add_reading_guide(document: Document) -> None:
    document.add_heading(humanize("Laesevejledning"), level=1)
    document.add_paragraph(
        humanize(
            "Manualen er lavet som et praktisk opslagsvaerk. Hver funktion er beskrevet kort, tydeligt og i et sprog, der tager udgangspunkt i arbejdsdagen frem for i den tekniske opbygning bag loesningen."
        )
    )
    add_bullets(
        document,
        [
            "Start i det relevante domaene og find derefter den konkrete funktion.",
            "Brug `Hvor findes den` som den hurtige vej ind i Business Central.",
            "Brug `Vigtige forhold` som kontrolpunkter foer eller under brug.",
            "Tekniske objektreferencer staar bagerst i hvert opslag og i appendikset.",
        ],
    )

    document.add_heading(humanize("Indholdsfortegnelse"), level=1)
    toc_paragraph = document.add_paragraph()
    add_field_code(toc_paragraph, r'TOC \o "1-3" \h \z \u')


def add_domain(document: Document, domain: dict[str, object]) -> None:
    document.add_heading(humanize(str(domain["name"])), level=1)
    banner_path = create_domain_banner(document._manual_assets_dir, str(domain["name"]), str(domain["intro"]))
    banner_paragraph = document.add_paragraph()
    banner_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    banner_paragraph.add_run().add_picture(str(banner_path), width=Inches(6.6))
    document.add_paragraph(humanize(str(domain["intro"])))
    add_callout_table(
        document,
        [
            ("Det faar du i dette omraade", str(domain["intro"])),
            ("Tone i vejledningen", "Kort, praktisk og lavet til hverdagsbrug i Business Central."),
        ],
        fill="EEE7DD",
    )

    for entry in domain["entries"]:
        add_entry(document, entry)


def add_entry(document: Document, entry: dict[str, object]) -> None:
    document.add_heading(humanize(str(entry["title"])), level=2)
    add_callout_table(
        document,
        [
            ("Det hjaelper funktionen dig med", str(entry["purpose"])),
            ("Fordel for dig", "Du faar hurtigere overblik og mindre manuel opfoelgning i den daglige opgave."),
        ],
        fill="F4EFE8",
    )
    add_label_paragraph(document, "Du finder den her", str(entry["location"]))
    add_label_paragraph(document, "Typisk naar du bruger den", str(entry["when"]))

    document.add_paragraph(humanize("Saadan kommer du godt i gang:"), style="Manual Label")
    add_numbered(document, entry["steps"])

    document.add_paragraph(humanize("Godt at vide:"), style="Manual Label")
    add_bullets(document, entry["important"])

    if entry.get("related"):
        add_label_paragraph(document, "Relaterede funktioner", ", ".join(entry["related"]))
    if entry.get("objects"):
        add_label_paragraph(document, "Teknisk reference", ", ".join(entry["objects"]))


def add_appendix(document: Document, inventory: list[ObjectInfo]) -> None:
    document.add_heading(humanize("Appendiks - objektoverblik"), level=1)
    document.add_paragraph(
        humanize(
            "Oversigten nedenfor er genereret fra de aktive AL-objekter i projektbiblioteket og kan bruges som teknisk reference ved senere opdateringer af manualen."
        )
    )

    table = document.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = [humanize(value) for value in ["Type", "Id", "Navn", "Caption", "Usage/PageType", "Fil"]]
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header

    for info in inventory:
        if info.kind not in {"page", "pageextension", "report", "reportextension"}:
            continue
        row = table.add_row().cells
        row[0].text = info.kind
        row[1].text = str(info.object_id)
        row[2].text = info.name if not info.extends else f"{info.name} -> {info.extends}"
        row[3].text = info.caption
        row[4].text = " / ".join(part for part in [info.usage_category, info.page_type] if part)
        row[5].text = info.file_path

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8.5)


def build_document(repo_root: Path) -> Document:
    document = Document()
    document._manual_assets_dir = repo_root / "docs" / "generated_manual_assets"
    document.sections[0].top_margin = Inches(0.7)
    document.sections[0].bottom_margin = Inches(0.7)
    document.sections[0].left_margin = Inches(0.8)
    document.sections[0].right_margin = Inches(0.8)
    configure_styles(document)

    app_info = json.loads((repo_root / "app.json").read_text(encoding="utf-8"))
    inventory = collect_inventory(repo_root)

    add_cover(document, app_info, repo_root)
    add_reading_guide(document)

    for domain in MANUAL_STRUCTURE["domains"]:
        add_domain(document, domain)

    add_appendix(document, inventory)
    return document


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate the BC25 end-user manual.")
    parser.add_argument("--repo-root", required=True, type=Path)
    parser.add_argument("--output-docx", required=True, type=Path)
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    repo_root = args.repo_root.resolve()
    output_docx = args.output_docx.resolve()

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    document = build_document(repo_root)
    document.save(output_docx)
    print(f"Generated: {output_docx}")


if __name__ == "__main__":
    main()
